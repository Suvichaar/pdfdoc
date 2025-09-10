import io
import os
import json
import re
import base64
import hashlib
from typing import Optional, Any, Dict
from datetime import datetime
from pathlib import Path

import streamlit as st

# =========================
# PAGE SETUP
# =========================
st.set_page_config(page_title="PDF → DOCX Suvichaars", page_icon="📄", layout="wide")
st.title("📄 PDF → DOCX with Suvichaar Document Intelligence")
st.caption(
    "Upload a PDF → SuvichaarDI (prebuilt-read) extracts text → Download a .docx • "
    "Pricing: ₹3 per page (3 credits) • Per-user credits set by Admin (no reset on reload)"
)

# =========================
# PRICING / CONSTANTS
# =========================
PRICE_PER_PAGE_CREDITS = 3  # ₹3 == 3 credits
DEFAULT_START_CREDITS = 30_000

# =========================
# SECRETS / CONFIG
# =========================
def get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    try:
        return st.secrets[key]  # type: ignore[attr-defined]
    except Exception:
        return default

AZURE_DI_ENDPOINT = get_secret("AZURE_DI_ENDPOINT")
AZURE_DI_KEY = get_secret("AZURE_DI_KEY")

# --- AWS / S3 (silent uploads) ---
AWS_REGION = get_secret("AWS_REGION", "ap-south-1")
AWS_ACCESS_KEY_ID = get_secret("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = get_secret("AWS_SECRET_ACCESS_KEY")
S3_BUCKET = get_secret("S3_BUCKET", "suvichaarapp")
S3_PREFIX = get_secret("S3_PREFIX", "media/pdf2docx")  # no leading slash

# --- Admin bootstrap ---
ADMIN_EMAIL = get_secret("ADMIN_EMAIL")
ADMIN_PASSWORD = get_secret("ADMIN_PASSWORD")  # first-run bootstrap

# --- Admin Panel PIN (6 digits) ---
# --- Admin Panel PIN (6 digits) ---
# Read ONLY from env or secrets; never hard-code a default.
ADMIN_PANEL_PIN = (os.getenv("ADMIN_PANEL_PIN") or get_secret("ADMIN_PANEL_PIN") or "").strip()
if not re.fullmatch(r"\d{6}", ADMIN_PANEL_PIN):
    st.error("ADMIN_PANEL_PIN missing/invalid. Set a 6-digit PIN via env var or .streamlit/secrets.toml and restart.")
    st.stop()


# =========================
# SDK IMPORTS
# =========================
try:
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential
except Exception:
    DocumentIntelligenceClient = None
    AzureKeyCredential = None

try:
    from azure.ai.documentintelligence.models import AnalyzeDocumentRequest  # type: ignore
except Exception:
    AnalyzeDocumentRequest = None  # type: ignore

from docx import Document
from docx.shared import Pt

# =========================
# USERS STORE (auth + wallet)
# =========================
USERS_STORE_PATH = Path("./users_store.json")
APP_SALT = b"SuvichaarDI_v1"  # app-level salt for PBKDF2

def _hash_pw(password: str, salt: bytes) -> str:
    h = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 120_000)
    return base64.b64encode(h).decode("utf-8")

def _set_pw(password: str) -> str:
    return _hash_pw(password, APP_SALT)

DEFAULT_USERS_DB = {"users": {}}  # email -> record

def load_users() -> Dict[str, Any]:
    if USERS_STORE_PATH.exists():
        try:
            with open(USERS_STORE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return DEFAULT_USERS_DB.copy()
    return DEFAULT_USERS_DB.copy()

def save_users(data: Dict[str, Any]) -> None:
    tmp = USERS_STORE_PATH.with_suffix(".tmp")
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, USERS_STORE_PATH)

# session bootstrap
if "users_db" not in st.session_state:
    st.session_state.users_db = load_users()
if "current_user" not in st.session_state:
    st.session_state.current_user = None
if "auth_view" not in st.session_state:
    st.session_state.auth_view = "login"
if "admin_panel_unlocked" not in st.session_state:
    st.session_state.admin_panel_unlocked = False  # gate admin panel with 6-digit PIN

# First-run: ensure admin exists
if ADMIN_EMAIL not in st.session_state.users_db["users"]:
    st.session_state.users_db["users"][ADMIN_EMAIL] = {
        "email": ADMIN_EMAIL,
        "name": "Admin",
        "tenant_id": "default-tenant",
        "profile_id": "admin-profile",
        "password_hash": _set_pw(ADMIN_PASSWORD),
        "force_pw_change": False,
        "is_admin": True,
        "start_credits": DEFAULT_START_CREDITS,
        "credits": DEFAULT_START_CREDITS,
        "ledger": [],                # [{file, pages, credits, ts}]
        "charged_docs": {},          # file_hash -> txn
        "last_txn": None,            # {file, pages, cost, ts}
        "last_s3_keys": [],          # [{type, key, ts}]
    }
    save_users(st.session_state.users_db)

# =========================
# S3 HELPERS (silent uploads)
# =========================
import boto3

def _sanitize_filename(name: str) -> str:
    base = name.strip().replace(" ", "_")
    return re.sub(r"[^A-Za-z0-9._-]+", "", base) or "file"

@st.cache_resource(show_spinner=False)
def _get_s3_client():
    if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
        return boto3.client(
            "s3",
            region_name=AWS_REGION,
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        )
    return boto3.client("s3", region_name=AWS_REGION)

def _build_object_key(prefix: str, kind: str, tenant_id: str, email: str, fid: str, filename: str, ext: str) -> str:
    safe = _sanitize_filename(filename.rsplit(".", 1)[0])
    today = datetime.now().strftime("%Y/%m/%d")
    email_key = email.replace("@", "_")
    return f"{(prefix or 'media/pdf2docx').rstrip('/')}/{kind}/{tenant_id}/{email_key}/{today}/{fid[:12]}-{safe}.{ext.lstrip('.')}"

def _put_bytes_to_s3(key: str, data: bytes, content_type: str) -> None:
    extra = {"ContentType": content_type}
    # extra["ServerSideEncryption"] = "AES256"  # uncomment if you want SSE-S3 explicitly
    _get_s3_client().put_object(Bucket=S3_BUCKET, Key=key, Body=data, **extra)

def silent_upload_pdf(fid: str, filename: str, pdf_bytes: bytes, tenant_id: str, email: str):
    try:
        key = _build_object_key(S3_PREFIX, "uploads", tenant_id, email, fid, filename, "pdf")
        _put_bytes_to_s3(key, pdf_bytes, "application/pdf")
        rec = get_user_rec()
        (rec.setdefault("last_s3_keys", [])).append({"type": "pdf", "key": key, "ts": datetime.now().isoformat()})
        save_user_rec(rec)
    except Exception:
        pass  # silent by design

def silent_upload_docx(fid: str, filename: str, docx_bytes: bytes, tenant_id: str, email: str):
    try:
        key = _build_object_key(S3_PREFIX, "outputs", tenant_id, email, fid, filename, "docx")
        _put_bytes_to_s3(key, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        rec = get_user_rec()
        (rec.setdefault("last_s3_keys", [])).append({"type": "docx", "key": key, "ts": datetime.now().isoformat()})
        save_user_rec(rec)
    except Exception:
        pass  # silent by design

# =========================
# AUTH UI
# =========================
def ui_login():
    st.subheader("Login")
    email = st.text_input("Email", key="auth_email")
    pw = st.text_input("Password", type="password", key="auth_pw")

    if st.button("Sign in", key="auth_signin_btn", use_container_width=True):
        rec = st.session_state.users_db["users"].get(email)
        if not rec:
            st.error("Invalid email or password.")
            return

        # If account flagged for first-time reset, only accept temporary password and redirect to Reset
        if rec.get("force_pw_change"):
            temp_hash = rec.get("temp_pw_hash")
            if temp_hash and _set_pw(pw) == temp_hash:
                st.session_state.auth_view = "reset"
                st.session_state.reset_email_prefill = email
                st.info("First-time login detected. Please set a new password to continue.")
                return
            else:
                st.error("This account requires a password reset. Use your temporary password to proceed.")
                return

        # Normal login (no reset required)
        if _set_pw(pw) == rec.get("password_hash"):
            st.session_state.current_user = rec
            st.success(f"Welcome {rec.get('name') or rec['email']}!")
        else:
            st.error("Invalid email or password.")

def ui_reset_password():
    st.subheader("Reset Password")
    st.caption("Use the temporary password you received from Admin once, then set a new password.")
    email = st.text_input("Email", value=st.session_state.get("reset_email_prefill",""), key="reset_email")
    temp_pw = st.text_input("Temporary Password", type="password", key="reset_temp")
    new_pw = st.text_input("New Password", type="password", key="reset_new")
    new_pw2 = st.text_input("Re-enter New Password", type="password", key="reset_new2")
    if st.button("Reset Password", key="reset_btn", use_container_width=True):
        rec = st.session_state.users_db["users"].get(email)
        if not rec:
            st.error("Account not found.")
            return
        if _set_pw(temp_pw) != rec.get("temp_pw_hash"):
            st.error("Temporary password incorrect.")
            return
        if not new_pw or new_pw != new_pw2:
            st.error("New passwords do not match.")
            return
        rec["password_hash"] = _set_pw(new_pw)
        rec["force_pw_change"] = False
        rec.pop("temp_pw_hash", None)
        save_user_rec(rec)
        st.success("Password updated. Please login.")
        st.session_state.auth_view = "login"

with st.sidebar:
    st.markdown("### Navigation")
    nav_choice = st.radio(
        label="",
        options=["Login", "Reset Password"],
        index=0 if st.session_state.auth_view == "login" else 1,
        key="auth_nav",
    )
    st.session_state.auth_view = "login" if nav_choice == "Login" else "reset"

# Gate until login succeeds
if st.session_state.current_user is None:
    if st.session_state.auth_view == "login":
        ui_login()
    else:
        ui_reset_password()
    st.stop()

# =========================
# PER-USER HELPERS
# =========================
def get_user_rec() -> Dict[str, Any]:
    return st.session_state.current_user

def save_user_rec(rec: Dict[str, Any]) -> None:
    db = st.session_state.users_db
    db["users"][rec["email"]] = rec
    save_users(db)
    st.session_state.current_user = rec  # keep session in sync

def file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def charge_user_for_pages(rec: Dict[str, Any], fid: str, pages: int, filename: str) -> int:
    """Deduct credits once per user+file_hash; persist to ledger & last_txn."""
    pages = max(1, int(pages))
    cost = pages * PRICE_PER_PAGE_CREDITS

    charged_docs = rec.setdefault("charged_docs", {})
    if fid in charged_docs:
        # Already billed for this file hash
        return 0

    if int(rec.get("credits", 0)) < cost:
        raise RuntimeError(f"Insufficient credits: need {cost}, have {rec.get('credits',0)}.")

    rec["credits"] = int(rec.get("credits", 0)) - cost
    txn = {"file": filename, "pages": pages, "cost": cost, "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
    rec["last_txn"] = txn
    (rec.setdefault("ledger", [])).append({"file": filename, "pages": pages, "credits": cost, "ts": txn["ts"]})
    charged_docs[fid] = txn
    save_user_rec(rec)
    return cost

# =========================
# SIDEBAR: PROFILE + CREDITS + ADMIN (with 6-digit PIN gate)
# =========================
with st.sidebar:
    u = get_user_rec()

    # Profile card
    st.markdown(
        f"""
        <div style="background:#111827;border:1px solid #374151;border-radius:12px;padding:12px;margin-bottom:10px;">
          <div style="font-size:16px;font-weight:600;">👤 Profile</div>
          <div style="font-size:13px;opacity:.9;margin-top:6px;">
            <div><b>Email:</b> {u['email']}</div>
            <div><b>Tenant ID:</b> {u.get('tenant_id','-')}</div>
            <div><b>Profile ID:</b> {u.get('profile_id','-')}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Credits meter
    st.subheader("💳 Credits")
    start_cap = max(int(u.get("start_credits", DEFAULT_START_CREDITS)), int(u.get("credits", 0)))
    pct = max(0.0, min(float(u.get("credits", 0)) / float(start_cap or 1), 1.0))
    st.progress(pct, text=f"Balance: {int(u.get('credits',0))} credits")
    st.caption("Pricing: 3 credits (₹3) per page • Set by Admin")

    # Last transaction (optional)
    txn = u.get("last_txn")
    if txn:
        st.markdown(
            f"""
            <div style="background:#f5f8ff;padding:12px;border-radius:10px;border:1px solid #d1e3ff;margin-top:12px;">
              <div style="font-weight:600;color:#1f4396;margin-bottom:6px;">🧾 Last Transaction</div>
              <div style="font-size:13px;line-height:1.4;">
                <div><b>File:</b> {txn['file']}</div>
                <div><b>Pages:</b> {txn['pages']}</div>
                <div><b>Credits:</b> {txn['cost']} (₹{txn['cost']})</div>
                <div style="color:#666;"><b>Time:</b> {txn.get('ts','')}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Admin panel (only if user is admin) + requires 6-digit PIN
    if u.get("is_admin"):
        with st.expander("🔐 Admin Panel", expanded=False):
            if not st.session_state.admin_panel_unlocked:
                st.info("Enter the 6-digit Admin PIN to unlock the panel.")
                pin_in = st.text_input("Admin PIN (6 digits)", type="password", max_chars=6, key="admin_pin_input")
                if st.button("Unlock Admin Panel", key="admin_pin_btn"):
                    if re.fullmatch(r"\d{6}", str(pin_in or "")) and str(pin_in) == ADMIN_PANEL_PIN:
                        st.session_state.admin_panel_unlocked = True
                        st.success("Admin Panel unlocked.")
                    else:
                        st.error("Invalid PIN. Please try again.")
            else:
                # Lock button
                if st.button("🔒 Lock Admin Panel", key="admin_pin_lock"):
                    st.session_state.admin_panel_unlocked = False

                st.markdown("**Create / Edit User**")
                a_email = st.text_input("User Email", key="a_email")
                a_name = st.text_input("Name", key="a_name")
                a_tenant = st.text_input("Tenant ID", key="a_tenant")
                a_profile = st.text_input("Profile ID", key="a_profile")
                a_start = st.number_input("Start Credits", min_value=0, value=DEFAULT_START_CREDITS, step=100, key="a_start")
                a_credits = st.number_input("Current Credits", min_value=0, value=DEFAULT_START_CREDITS, step=100, key="a_credits")
                a_temp_pw = st.text_input("Temporary Password (for new/reset)", type="password", key="a_temp")

                if st.button("Save User", key="a_save"):
                    if not a_email:
                        st.error("Email required.")
                    else:
                        db = st.session_state.users_db
                        rec = db["users"].get(a_email, {})
                        rec.update({
                            "email": a_email,
                            "name": a_name or rec.get("name") or "",
                            "tenant_id": a_tenant or rec.get("tenant_id") or "",
                            "profile_id": a_profile or rec.get("profile_id") or "",
                            "is_admin": rec.get("is_admin", False),
                            "start_credits": int(a_start),
                            "credits": int(a_credits),
                            "ledger": rec.get("ledger", []),
                            "charged_docs": rec.get("charged_docs", {}),
                            "last_txn": rec.get("last_txn", None),
                            "last_s3_keys": rec.get("last_s3_keys", []),
                        })
                        if a_temp_pw:
                            rec["temp_pw_hash"] = _set_pw(a_temp_pw)
                            rec["force_pw_change"] = True
                        else:
                            rec["force_pw_change"] = rec.get("force_pw_change", False)

                        db["users"][a_email] = rec
                        save_users(db)
                        st.success("User saved / updated.")

                st.markdown("---")
                st.markdown("**Top-up Credits**")
                top_email = st.text_input("Email to top-up", key="top_email")
                top_amt = st.number_input("Amount", min_value=1, value=100, step=50, key="top_amt")
                if st.button("Top-up", key="top_btn"):
                    db = st.session_state.users_db
                    rec = db["users"].get(top_email)
                    if not rec:
                        st.error("User not found.")
                    else:
                        rec["credits"] = int(rec.get("credits", 0)) + int(top_amt)
                        save_users(db)
                        st.success(f"Topped up {top_amt} credits.")

                st.markdown("---")
                st.markdown("**Grant/Revoke Admin**")
                adm_email = st.text_input("Email", key="adm_email")
                make_admin = st.checkbox("Is Admin?", value=False, key="adm_flag")
                if st.button("Update Admin Flag", key="adm_btn"):
                    db = st.session_state.users_db
                    rec = db["users"].get(adm_email)
                    if not rec:
                        st.error("User not found.")
                    else:
                        rec["is_admin"] = bool(make_admin)
                        save_users(db)
                        st.success("Updated.")

                # =========================
                # NEW: Set Tenant/Profile for any user (after PIN unlock)
                # =========================
                st.markdown("---")
                st.markdown("**Set Tenant ID / Profile ID**")

                users_map = st.session_state.users_db.get("users", {})
                user_emails = sorted(users_map.keys())

                sel_email = st.selectbox("Select user", options=user_emails, key="tenant_profile_sel_email")

                if sel_email:
                    target = users_map.get(sel_email, {})
                    cur_tenant  = target.get("tenant_id", "")
                    cur_profile = target.get("profile_id", "")

                    new_tenant  = st.text_input("Tenant ID",  value=cur_tenant,  key="tenant_profile_new_tenant")
                    new_profile = st.text_input("Profile ID", value=cur_profile, key="tenant_profile_new_profile")

                    if st.button("Save Tenant/Profile", key="tenant_profile_save_btn"):
                        target["tenant_id"]  = (new_tenant or "").strip()
                        target["profile_id"] = (new_profile or "").strip()

                        db = st.session_state.users_db
                        db["users"][sel_email] = target
                        save_users(db)

                        # keep session in sync if current user was updated
                        if st.session_state.current_user["email"] == sel_email:
                            st.session_state.current_user = target

                        st.success(f"Updated tenant/profile for {sel_email}.")

# =========================
# SETTINGS (single expander)
# =========================
with st.expander("⚙️ Settings", expanded=False):
    add_page_breaks = st.checkbox("Insert page breaks between PDF pages", value=True, key="opt_page_breaks")
    include_confidence = st.checkbox("Append line confidence (debug)", value=False, key="opt_conf")

# If secrets not configured, allow input for this run
if not AZURE_DI_ENDPOINT or not AZURE_DI_KEY:
    st.info("SuvichaarDI endpoint/key not found in st.secrets. Enter them for this session.")
    AZURE_DI_ENDPOINT = st.text_input(
        "AZURE_DI_ENDPOINT",
        AZURE_DI_ENDPOINT or "",
        placeholder="https://<resourcename>.cognitiveservices.azure.com/",
        key="endpoint_input",
    )
    AZURE_DI_KEY = st.text_input(
        "AZURE_DI_KEY",
        AZURE_DI_KEY or "",
        type="password",
        key="key_input",
    )

# =========================
# HELPERS (Azure DI + DOCX)
# =========================
@st.cache_resource(show_spinner=False)
def make_client(endpoint: str, key: str):
    if DocumentIntelligenceClient is None or AzureKeyCredential is None:
        raise RuntimeError("SuvichaarSDK not installed. Run: pip install azure-ai-documentintelligence python-docx")
    if not endpoint or not key:
        raise RuntimeError("Missing SuvichaarDI endpoint or key.")
    return DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))

def analyze_pdf_bytes(client: Any, pdf_bytes: bytes):
    last_err = None
    try:
        poller = client.begin_analyze_document(
            model_id="prebuilt-read",
            document=pdf_bytes,
            content_type="application/pdf",
        )
        return poller.result()
    except Exception as e:
        last_err = e
    try:
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        poller = client.begin_analyze_document(
            "prebuilt-read",
            body={"base64Source": b64},
        )
        return poller.result()
    except Exception as e:
        last_err = e
    try:
        if AnalyzeDocumentRequest is not None:
            req = AnalyzeDocumentRequest(bytes_source=pdf_bytes)  # type: ignore
            poller = client.begin_analyze_document(
                model_id="prebuilt-read",
                body=req,
                content_type="application/pdf",
            )
            return poller.result()
    except Exception as e:
        last_err = e
    raise last_err

def result_to_docx_bytes(result, insert_page_breaks: bool = True, show_conf: bool = False) -> bytes:
    """Convert DI 'prebuilt-read' result into a .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if not getattr(result, "pages", None):
        doc.add_paragraph(getattr(result, "content", "") or "No content found.")
    else:
        for idx, page in enumerate(result.pages):
            doc.add_heading(f"Page {idx+1}", level=2)
            if getattr(page, "lines", None):
                for ln in page.lines:
                    text = ln.content or ""
                    if show_conf and hasattr(ln, "spans") and ln.spans:
                        try:
                            confs = [getattr(s, "confidence", None) for s in ln.spans if getattr(s, "confidence", None) is not None]
                            if confs:
                                text += f"  [conf~{sum(confs)/len(confs):.2f}]"
                        except Exception:
                            pass
                    if text.strip():
                        doc.add_paragraph(text)
            else:
                paras = []
                for p in getattr(result, "paragraphs", []) or []:
                    if getattr(p, "spans", None):
                        if any(getattr(sp, "offset", None) is not None for sp in p.spans):
                            paras.append(p.content)
                if paras:
                    for p in paras:
                        doc.add_paragraph(p)
                else:
                    doc.add_paragraph(getattr(result, "content", "") or "")

            if insert_page_breaks and idx < len(result.pages) - 1:
                doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# =========================
# MAIN FLOW
# =========================
uploaded = st.file_uploader("Upload a PDF", type=["pdf"], accept_multiple_files=False, key="pdf_uploader_main")

if uploaded is not None:
    if not uploaded.name.lower().endswith(".pdf"):
        st.error("Please upload a PDF file.")
    else:
        try:
            client = make_client(AZURE_DI_ENDPOINT or "", AZURE_DI_KEY or "")
        except Exception as e:
            st.error(f"Failed to create SuvichaarDI client: {e}")
            st.stop()

        pdf_bytes = uploaded.read()
        if not pdf_bytes:
            st.error("Uploaded file is empty. Please re-upload the PDF.")
            st.stop()

        fid = file_hash(pdf_bytes)

        # Silent S3 upload (PDF) before analysis
        u = get_user_rec()
        silent_upload_pdf(fid, uploaded.name, pdf_bytes, u.get("tenant_id", "default"), u["email"])

        with st.spinner("Analyzing with SuvichaarDocument Intelligence (prebuilt-read)..."):
            try:
                result = analyze_pdf_bytes(client, pdf_bytes)
            except Exception as e:
                st.error(f"SuvichaarDI analyze failed: {e}")
                st.stop()

        pages = len(getattr(result, "pages", []) or [])
        if pages <= 0:
            pages = 1
        st.success(f"Extracted text from **{pages} page(s)**.")

        # Billing (once per user+file hash)
        try:
            charged = charge_user_for_pages(u, fid, pages, uploaded.name)
            if charged > 0:
                st.toast(f"Charged {charged} credits for {pages} page(s).", icon="✅")
        except RuntimeError as e:
            st.error(str(e))
            st.stop()

        with st.spinner("Building DOCX..."):
            try:
                docx_bytes = result_to_docx_bytes(
                    result,
                    insert_page_breaks=st.session_state.get("opt_page_breaks", True),
                    show_conf=st.session_state.get("opt_conf", False)
                )
            except Exception as e:
                st.error(f"Failed to create DOCX: {e}")
                st.stop()

        # Silent S3 upload (DOCX) after build
        docx_filename = (uploaded.name.rsplit(".", 1)[0] + ".docx")
        silent_upload_docx(fid, docx_filename, docx_bytes, u.get("tenant_id", "default"), u["email"])

        # Local download only (no S3/CDN links shown)
        st.download_button(
            label="⬇️ Download .docx",
            data=docx_bytes,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_btn"
        )
else:
    st.info("Upload a PDF to begin.")

# =========================
# FOOTER
# =========================
st.caption(
    "Per-user credits persist across reloads (stored server-side). Admin creates users, sets tenant/profile & credits, "
    f"and can top-up anytime. Pricing: {PRICE_PER_PAGE_CREDITS} credits (₹{PRICE_PER_PAGE_CREDITS}) per page."
)
